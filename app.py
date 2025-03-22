import streamlit as st
import docx
import re
import io
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from wordcloud import WordCloud
import matplotlib.pyplot as plt

# Function to read DOCX file
def read_docx(file):
    doc = docx.Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return doc, text

# Function to read bulk processing words from TXT file
def read_txt(file):
    return file.read().decode("utf-8").splitlines()

# Function to process DOCX modifications
def process_text(doc, remove_words, replace_dict, style_dict, selected_pages, regex_mode):
    for i, para in enumerate(doc.paragraphs):
        if selected_pages and i not in selected_pages:
            continue

        for word in remove_words:
            para.text = para.text.replace(word, "")

        for old, new in replace_dict.items():
            if regex_mode:
                para.text = re.sub(old, new, para.text)
            else:
                para.text = para.text.replace(old, new)

        for word, style in style_dict.items():
            for run in para.runs:
                if word in run.text:
                    run.bold = style.get("bold", False)
                    run.italic = style.get("italic", False)
                    run.font.size = Pt(style.get("size", 12))
                    if "color" in style:
                        run.font.color.rgb = RGBColor(*style["color"])
                    if "bg_color" in style:
                        run.font.highlight_color = style["bg_color"]
    
    return doc

# Function to modify tables in DOCX
def process_tables(doc, replace_dict):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old, new in replace_dict.items():
                    cell.text = re.sub(old, new, cell.text) if regex_mode else cell.text.replace(old, new)

# Function to generate a word cloud
def generate_wordcloud(text):
    wordcloud = WordCloud(width=600, height=300, background_color="white").generate(text)
    fig, ax = plt.subplots()
    ax.imshow(wordcloud, interpolation="bilinear")
    ax.axis("off")
    return fig

# Streamlit UI
st.title("📄 Advanced DOCX Editor with Full Customization")

# Upload DOCX file
uploaded_file = st.file_uploader("Upload a DOCX file", type="docx")

if uploaded_file:
    doc, original_text = read_docx(uploaded_file)
    st.text_area("Original Text", original_text, height=200)

    # Word Cloud Visualization
    if st.checkbox("Show Word Cloud"):
        st.pyplot(generate_wordcloud(original_text))

    # Bulk Processing via TXT Upload
    txt_file = st.file_uploader("Upload a .txt file for bulk processing (optional)", type="txt")
    bulk_words = read_txt(txt_file) if txt_file else []

    # User Inputs
    remove_words = st.text_area("Words to Remove (comma-separated)").split(",") + bulk_words

    replace_text = st.text_area("Find & Replace (format: old=new, comma-separated)")
    replace_dict = {pair.split("=")[0].strip(): pair.split("=")[1].strip() for pair in replace_text.split(",") if "=" in pair}

    style_text = st.text_area("Words to Style (format: word:bold:yes,size:14,color:255,0,0)")
    style_dict = {}
    for line in style_text.split("\n"):
        parts = line.split(":")
        if len(parts) >= 2:
            word = parts[0].strip()
            style_dict[word] = {
                "bold": "bold:yes" in parts,
                "italic": "italic:yes" in parts,
                "size": int(parts[parts.index("size") + 1]) if "size" in parts else 12,
                "color": tuple(map(int, parts[parts.index("color") + 1].split(","))) if "color" in parts else (0, 0, 0)
            }

    regex_mode = st.checkbox("Enable Regex Mode for Find & Replace")

    # Select Pages
    num_paragraphs = len(doc.paragraphs)
    selected_pages = st.multiselect("Select Pages to Modify", range(num_paragraphs), default=list(range(num_paragraphs)))

    # Preview Before Applying
    preview_text = original_text
    for word in remove_words:
        preview_text = preview_text.replace(word, "")

    for old, new in replace_dict.items():
        preview_text = re.sub(old, new, preview_text) if regex_mode else preview_text.replace(old, new)

    st.text_area("Preview Modified Text", preview_text, height=200)

    if st.button("Apply Changes"):
        modified_doc = process_text(doc, remove_words, replace_dict, style_dict, selected_pages, regex_mode)
        process_tables(modified_doc, replace_dict)

        # Save modified DOCX to buffer
        buffer = io.BytesIO()
        modified_doc.save(buffer)
        buffer.seek(0)

        # Download button
        st.download_button("Download Modified DOCX", buffer, "modified.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
